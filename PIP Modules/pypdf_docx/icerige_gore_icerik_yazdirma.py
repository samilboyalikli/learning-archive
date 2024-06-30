
import os
# silme işlemi için os modülünü içeri aktarıyorum
import time
# time.sleep() fonksiyonunu kullanmak için time modülünü içeri aktarıyorum
from docx import Document
from PyPDF2 import PdfReader

ss = 1

def belge_adi():
    # belge adini tanımlıyoruz
    isim = "yeni_belge"
    uzanti = ".docx"
    return isim + str(ss) + uzanti
def adres_adi():
    # belgenin bulunacağı adresi tanımlıyoruz
    isim = "C:\\Users\\bemco\\PycharmProjects\\pythonProject\\yeni_belge"
    uzanti = ".docx"
    return isim + str(ss) + uzanti
def cikti_adi():
    # en son elde edeceğimiz word dosyasının adını tanımlıyoruz
    isim = "yepyenibelge"
    uzanti = ".docx"
    return isim + str(ss) + uzanti

# işimizi kolaylaştıracak bir takım kısayollar oluşturuyoruz
belge = belge_adi()
adres = adres_adi()
cikti = cikti_adi()

# PDF'yi aç ve belge oluştur
reader = PdfReader("sozluk.pdf")
page = reader.pages[ss]
sonuc = page.extract_text()
belge_doc = Document()
belge_doc.add_paragraph(sonuc)
belge_doc.save(belge)

# belge kaydedildi şimdi programımıza biraz zaman tanıyacağız
time.sleep(5)  # 5 saniye belgenin kaydedilmiş olması için gayet yeterli bir süre

# belgeyi açıyoruz
doc = Document(adres)

# burada metin için gerekli editleri yapıyoruz
for paragraph in doc.paragraphs:
    # paragraftaki metni satırlara bölelim
    lines = paragraph.text.splitlines()

    # satırları düzenleyelim: "*" ile başlayan satırları çıkaralım
    lines = [line for line in lines if not line.startswith("*")]

    # paragrafın içeriğini varsayılan şekle getiriyoruz
    paragraph.clear()

    # düzenlenmiş satırları paragrafın yerine ekliyoruz
    for line in lines:
        paragraph.add_run(line + "\n")

# dosyayı kaydediyoruz
doc.save(cikti)

# eski belgeyi silmek için
os.remove(belge)
