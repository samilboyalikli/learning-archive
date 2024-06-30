import os
import time
from docx import Document
from PyPDF2 import PdfReader

# ellibin tane kod bloğu kopyalamamıza gerek kalmadan istediğimiz sayıda
# sayfayı pdf'ten word'e çevirip işlem yapabileceğimiz bir program
# python'da bir fonksiyon keşfettim, 400 satırda yaptığım şeyi tek kelimeyle yapabiliriz dsfgdsfds

def belge_adi(s):
    isim = "yeni_belge"
    uzanti = ".docx"
    return isim + str(s) + uzanti
def adres_adi(s):
    isim = "C:\\Users\\bemco\\PycharmProjects\\pythonProject\\yeni_belge"
    uzanti = ".docx"
    return isim + str(s) + uzanti
def cikti_adi(s):
    isim = "yepyenibelge"
    uzanti = ".docx"
    return isim + str(s) + uzanti
def pdften_text_alma(s):
    reader = PdfReader("sozluk.pdf")
    page = reader.pages[s]
    return page.extract_text()

def textle_word_dosyasi_yapma(text, belge):
    belge_doc = Document()
    belge_doc.add_paragraph(text)
    belge_doc.save(belge)
def paragraftaki_islemler(adres, cikti):
    doc = Document(adres)
    for paragraph in doc.paragraphs:
        lines = paragraph.text.splitlines()
        lines = [line for line in lines if not line.startswith("*")]
        paragraph.clear()
        for line in lines:
            paragraph.add_run(line + "\n")
    doc.save(cikti)
def belgelerdeki_islemler(num_of_documents):
    for s in range(num_of_documents):
        # bizi yaklaşık 400 satır koddan kurtaran fonksiyon: range()
        # bu fonksiyon sıfırdan başlayıp parantezin içine yazdığımız
        # sayıya kadar belgede istediğimiz işlemi yapıyor
        belge = belge_adi(s)
        adres = adres_adi(s)
        cikti = cikti_adi(s)

        text = pdften_text_alma(s)
        textle_word_dosyasi_yapma(text, belge)

        time.sleep(5)

        paragraftaki_islemler(adres, cikti)
        os.remove(belge)
        time.sleep(5)

    birlesik_belge = Document()
    for s in range(num_of_documents):
        belge = Document(cikti_adi(s))
        for paragraf in belge.paragraphs:
            yeni_paragraf = birlesik_belge.add_paragraph(paragraf.text)

    birlesik_belge.save("anabelge.docx")

# num_of_documents değişkeni ile kaç belge olacağını belirliyoruz
# 0 tabanlı indexleme kullanıyor program, dolayısıyla işlem yapacağınız-
# son sayfadan bir fazlasını girmeniz gerek
num_of_documents = 21
belgelerdeki_islemler(num_of_documents)
