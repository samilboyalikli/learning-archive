import os
import time
from docx import Document
from PyPDF2 import PdfReader

hangi_sayfaya_kadar = input("Metni hangi sayfaya kadar dönüştüreceğinizi giriniz: ")
hangi_sayfaya_kadar = int(hangi_sayfaya_kadar)+1
num_of_documents = hangi_sayfaya_kadar

def belge_adi(s):
    isim = "yeni_belge"
    uzanti = ".docx"
    return isim + str(s) + uzanti
def adres_adi(s):
    isim = "C:\\Users\\bemco\\PycharmProjects\\pythonProject\\yeni_belge"
    uzanti = ".docx"
    return isim + str(s) + uzanti
def cikti_adi(s):
    isim = "yepyenibelge"
    uzanti = ".docx"
    return isim + str(s) + uzanti
def pdften_text_alma(s):
    reader = PdfReader("sozluk.pdf")
    page = reader.pages[s]
    return page.extract_text()
def textle_word_dosyasi_yapma(text, belge):
    belge_doc = Document()
    belge_doc.add_paragraph(text)
    belge_doc.save(belge)
def paragraftaki_islemler(adres, cikti):
    doc = Document(adres)
    for paragraph in doc.paragraphs:
        lines = paragraph.text.splitlines()
        lines = [line for line in lines if not line.startswith("*")]
        paragraph.clear()
        for line in lines:
            paragraph.add_run(line + "\n")
    doc.save(cikti)
def belgelerdeki_islemler(num_of_documents):
    for s in range(num_of_documents):
        belge = belge_adi(s)
        adres = adres_adi(s)
        cikti = cikti_adi(s)

        text = pdften_text_alma(s)
        textle_word_dosyasi_yapma(text, belge)

        time.sleep(3)

        paragraftaki_islemler(adres, cikti)
        os.remove(belge)

    birlesik_belge = Document()
    for s in range(num_of_documents):
        belge = Document(cikti_adi(s))
        for paragraf in belge.paragraphs:
            yeni_paragraf = birlesik_belge.add_paragraph(paragraf.text)

    birlesik_belge.save("anabelge.docx")
belgelerdeki_islemler(num_of_documents)
def eski_dosyalarin_silinmesi(num_of_documents):
    for s in range(num_of_documents):
        # num_of_documents değerine kadar (ki bu inputta user'ın gireceği değer) tüm s değerlerinde dön
        dosya_yolu = f"yepyenibelge{s}.docx"
        # dosya yolları bu şekilde olacak
        os.remove(dosya_yolu)
#       belirttiğim yoldaki dosyaları sil
eski_dosyalarin_silinmesi(num_of_documents)
