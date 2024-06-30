# düzenlemeden word dosyasına aktaracağımız alternatif kod bloğu (we = without edit)
from pypdf import PdfReader
from docx import Document
# docx kütüphanesinden Document modülünü içeri aktarıyoruz
# docx indirek için:
# pip install python-docx

ss = 3
# sayfanın numarasını (0 tabanı indeks kullanarak) tanımlıyoruz
reader = PdfReader("sozluk.pdf")
page = reader.pages[ss]
sonuc = page.extract_text()
# çıktıyı sonuc adıyla belirtiyoruz
doc = Document()
# Document() fonksiyonunu "doc" olarak kısaltıyoruz
doc.add_paragraph(sonuc)
# dosyaya paragraf olarak sonuc adlı veriyi ekle
doc.save("yeni_belge3.docx")
# dosyayı belirttiğim isimle kaydet
