# pdf dosyanızın tek bir sayfa sayısını giriyorsunuz program sonraki 10 sayfada da işlem yapıyor (0 tabanlı indexlemeyle)
import os
import time
from docx import Document
from PyPDF2 import PdfReader
ss = 0
# sayfa sayısını buraya giriyoruz
# sayfa sayısının kullanılacağı fonksiyonlara direkt aşağıdaki sayfa numaralarını koydum
s1 = ss+1
s2 = s1+1
s3 = s2+1
s4 = s3+1
s5 = s4+1
s6 = s5+1
s7 = s6+1
s8 = s7+1
s9 = s8+1
s10 = s9+1
def belge_adi():
    isim = "yeni_belge"
    uzanti = ".docx"
    return isim + str(ss) + uzanti
def adres_adi():
    isim = "C:\\Users\\bemco\\PycharmProjects\\pythonProject\\yeni_belge"
    uzanti = ".docx"
    return isim + str(ss) + uzanti
def cikti_adi():
    isim = "yepyenibelge"
    uzanti = ".docx"
    return isim + str(ss) + uzanti

belge = belge_adi()
adres = adres_adi()
cikti = cikti_adi()
reader = PdfReader("sozluk.pdf")
page = reader.pages[ss]
sonuc = page.extract_text()
belge_doc = Document()
belge_doc.add_paragraph(sonuc)
belge_doc.save(belge)

time.sleep(5)

doc = Document(adres)
for paragraph in doc.paragraphs:
    lines = paragraph.text.splitlines()
    lines = [line for line in lines if not line.startswith("*")]
    paragraph.clear()
    for line in lines:
        paragraph.add_run(line + "\n")
doc.save(cikti)
os.remove(belge)
# --------------------------------------------------------------------------------
def belge_adi():
    isim = "yeni_belge"
    uzanti = ".docx"
    return isim + str(s1) + uzanti

def adres_adi():
    isim = "C:\\Users\\bemco\\PycharmProjects\\pythonProject\\yeni_belge"
    uzanti = ".docx"
    return isim + str(s1) + uzanti
def cikti_adi():
    isim = "yepyenibelge"
    uzanti = ".docx"
    return isim + str(s1) + uzanti

belge = belge_adi()
adres = adres_adi()
cikti = cikti_adi()
reader = PdfReader("sozluk.pdf")
page = reader.pages[s1]
sonuc = page.extract_text()
belge_doc = Document()
belge_doc.add_paragraph(sonuc)
belge_doc.save(belge)

time.sleep(5)

doc = Document(adres)
for paragraph in doc.paragraphs:
    lines = paragraph.text.splitlines()
    lines = [line for line in lines if not line.startswith("*")]
    paragraph.clear()
    for line in lines:
        paragraph.add_run(line + "\n")
doc.save(cikti)
os.remove(belge)
# ------------------------------------------------------------------
def belge_adi():
    isim = "yeni_belge"
    uzanti = ".docx"
    return isim + str(s2) + uzanti
def adres_adi():
    isim = "C:\\Users\\bemco\\PycharmProjects\\pythonProject\\yeni_belge"
    uzanti = ".docx"
    return isim + str(s2) + uzanti
def cikti_adi():
    isim = "yepyenibelge"
    uzanti = ".docx"
    return isim + str(s2) + uzanti

belge = belge_adi()
adres = adres_adi()
cikti = cikti_adi()
reader = PdfReader("sozluk.pdf")
page = reader.pages[s2]
sonuc = page.extract_text()
belge_doc = Document()
belge_doc.add_paragraph(sonuc)
belge_doc.save(belge)

time.sleep(5)

doc = Document(adres)
for paragraph in doc.paragraphs:
    lines = paragraph.text.splitlines()
    lines = [line for line in lines if not line.startswith("*")]
    paragraph.clear()
    for line in lines:
        paragraph.add_run(line + "\n")
doc.save(cikti)
os.remove(belge)
# ------------------------------------------------------------------
def belge_adi():
    isim = "yeni_belge"
    uzanti = ".docx"
    return isim + str(s3) + uzanti
def adres_adi():
    isim = "C:\\Users\\bemco\\PycharmProjects\\pythonProject\\yeni_belge"
    uzanti = ".docx"
    return isim + str(s3) + uzanti
def cikti_adi():
    isim = "yepyenibelge"
    uzanti = ".docx"
    return isim + str(s3) + uzanti

belge = belge_adi()
adres = adres_adi()
cikti = cikti_adi()
reader = PdfReader("sozluk.pdf")
page = reader.pages[s3]
sonuc = page.extract_text()
belge_doc = Document()
belge_doc.add_paragraph(sonuc)
belge_doc.save(belge)

time.sleep(5)

doc = Document(adres)
for paragraph in doc.paragraphs:
    lines = paragraph.text.splitlines()
    lines = [line for line in lines if not line.startswith("*")]
    paragraph.clear()
    for line in lines:
        paragraph.add_run(line + "\n")
doc.save(cikti)
os.remove(belge)
# ------------------------------------------------------------------
def belge_adi():
    isim = "yeni_belge"
    uzanti = ".docx"
    return isim + str(s4) + uzanti
def adres_adi():
    isim = "C:\\Users\\bemco\\PycharmProjects\\pythonProject\\yeni_belge"
    uzanti = ".docx"
    return isim + str(s4) + uzanti
def cikti_adi():
    isim = "yepyenibelge"
    uzanti = ".docx"
    return isim + str(s4) + uzanti

belge = belge_adi()
adres = adres_adi()
cikti = cikti_adi()
reader = PdfReader("sozluk.pdf")
page = reader.pages[s4]
sonuc = page.extract_text()
belge_doc = Document()
belge_doc.add_paragraph(sonuc)
belge_doc.save(belge)

time.sleep(5)

doc = Document(adres)
for paragraph in doc.paragraphs:
    lines = paragraph.text.splitlines()
    lines = [line for line in lines if not line.startswith("*")]
    paragraph.clear()
    for line in lines:
        paragraph.add_run(line + "\n")
doc.save(cikti)
os.remove(belge)
# ------------------------------------------------------------------
def belge_adi():
    isim = "yeni_belge"
    uzanti = ".docx"
    return isim + str(s5) + uzanti
def adres_adi():
    isim = "C:\\Users\\bemco\\PycharmProjects\\pythonProject\\yeni_belge"
    uzanti = ".docx"
    return isim + str(s5) + uzanti
def cikti_adi():
    isim = "yepyenibelge"
    uzanti = ".docx"
    return isim + str(s5) + uzanti

belge = belge_adi()
adres = adres_adi()
cikti = cikti_adi()
reader = PdfReader("sozluk.pdf")
page = reader.pages[s5]
sonuc = page.extract_text()
belge_doc = Document()
belge_doc.add_paragraph(sonuc)
belge_doc.save(belge)

time.sleep(5)

doc = Document(adres)
for paragraph in doc.paragraphs:
    lines = paragraph.text.splitlines()
    lines = [line for line in lines if not line.startswith("*")]
    paragraph.clear()
    for line in lines:
        paragraph.add_run(line + "\n")
doc.save(cikti)
os.remove(belge)
# ------------------------------------------------------------------
def belge_adi():
    isim = "yeni_belge"
    uzanti = ".docx"
    return isim + str(s6) + uzanti
def adres_adi():
    isim = "C:\\Users\\bemco\\PycharmProjects\\pythonProject\\yeni_belge"
    uzanti = ".docx"
    return isim + str(s6) + uzanti
def cikti_adi():
    isim = "yepyenibelge"
    uzanti = ".docx"
    return isim + str(s6) + uzanti

belge = belge_adi()
adres = adres_adi()
cikti = cikti_adi()
reader = PdfReader("sozluk.pdf")
page = reader.pages[s6]
sonuc = page.extract_text()
belge_doc = Document()
belge_doc.add_paragraph(sonuc)
belge_doc.save(belge)

time.sleep(5)

doc = Document(adres)
for paragraph in doc.paragraphs:
    lines = paragraph.text.splitlines()
    lines = [line for line in lines if not line.startswith("*")]
    paragraph.clear()
    for line in lines:
        paragraph.add_run(line + "\n")
doc.save(cikti)
os.remove(belge)
# ------------------------------------------------------------------
def belge_adi():
    isim = "yeni_belge"
    uzanti = ".docx"
    return isim + str(s7) + uzanti
def adres_adi():
    isim = "C:\\Users\\bemco\\PycharmProjects\\pythonProject\\yeni_belge"
    uzanti = ".docx"
    return isim + str(s7) + uzanti
def cikti_adi():
    isim = "yepyenibelge"
    uzanti = ".docx"
    return isim + str(s7) + uzanti

belge = belge_adi()
adres = adres_adi()
cikti = cikti_adi()
reader = PdfReader("sozluk.pdf")
page = reader.pages[s7]
sonuc = page.extract_text()
belge_doc = Document()
belge_doc.add_paragraph(sonuc)
belge_doc.save(belge)

time.sleep(5)

doc = Document(adres)
for paragraph in doc.paragraphs:
    lines = paragraph.text.splitlines()
    lines = [line for line in lines if not line.startswith("*")]
    paragraph.clear()
    for line in lines:
        paragraph.add_run(line + "\n")
doc.save(cikti)
os.remove(belge)
# ------------------------------------------------------------------
def belge_adi():
    isim = "yeni_belge"
    uzanti = ".docx"
    return isim + str(s8) + uzanti
def adres_adi():
    isim = "C:\\Users\\bemco\\PycharmProjects\\pythonProject\\yeni_belge"
    uzanti = ".docx"
    return isim + str(s8) + uzanti
def cikti_adi():
    isim = "yepyenibelge"
    uzanti = ".docx"
    return isim + str(s8) + uzanti

belge = belge_adi()
adres = adres_adi()
cikti = cikti_adi()
reader = PdfReader("sozluk.pdf")
page = reader.pages[s8]
sonuc = page.extract_text()
belge_doc = Document()
belge_doc.add_paragraph(sonuc)
belge_doc.save(belge)

time.sleep(5)

doc = Document(adres)
for paragraph in doc.paragraphs:
    lines = paragraph.text.splitlines()
    lines = [line for line in lines if not line.startswith("*")]
    paragraph.clear()
    for line in lines:
        paragraph.add_run(line + "\n")
doc.save(cikti)
os.remove(belge)
# ------------------------------------------------------------------
def belge_adi():
    isim = "yeni_belge"
    uzanti = ".docx"
    return isim + str(s9) + uzanti
def adres_adi():
    isim = "C:\\Users\\bemco\\PycharmProjects\\pythonProject\\yeni_belge"
    uzanti = ".docx"
    return isim + str(s9) + uzanti
def cikti_adi():
    isim = "yepyenibelge"
    uzanti = ".docx"
    return isim + str(s9) + uzanti

belge = belge_adi()
adres = adres_adi()
cikti = cikti_adi()
reader = PdfReader("sozluk.pdf")
page = reader.pages[s9]
sonuc = page.extract_text()
belge_doc = Document()
belge_doc.add_paragraph(sonuc)
belge_doc.save(belge)

time.sleep(5)

doc = Document(adres)
for paragraph in doc.paragraphs:
    lines = paragraph.text.splitlines()
    lines = [line for line in lines if not line.startswith("*")]
    paragraph.clear()
    for line in lines:
        paragraph.add_run(line + "\n")
doc.save(cikti)
os.remove(belge)
# ------------------------------------------------------------------
def belge_adi():
    isim = "yeni_belge"
    uzanti = ".docx"
    return isim + str(s10) + uzanti
def adres_adi():
    isim = "C:\\Users\\bemco\\PycharmProjects\\pythonProject\\yeni_belge"
    uzanti = ".docx"
    return isim + str(s10) + uzanti
def cikti_adi():
    isim = "yepyenibelge"
    uzanti = ".docx"
    return isim + str(s10) + uzanti

belge = belge_adi()
adres = adres_adi()
cikti = cikti_adi()
reader = PdfReader("sozluk.pdf")
page = reader.pages[s10]
sonuc = page.extract_text()
belge_doc = Document()
belge_doc.add_paragraph(sonuc)
belge_doc.save(belge)

time.sleep(5)

doc = Document(adres)
for paragraph in doc.paragraphs:
    lines = paragraph.text.splitlines()
    lines = [line for line in lines if not line.startswith("*")]
    paragraph.clear()
    for line in lines:
        paragraph.add_run(line + "\n")
doc.save(cikti)
os.remove(belge)
